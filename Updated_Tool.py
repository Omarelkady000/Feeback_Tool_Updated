import streamlit as st
import re
import math
import csv
import io
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# --- 1. CONFIGURATION DATA ---
XML_TIMEBASE_MAP = {
    "10.00 fps": 10, "12.00 fps": 12, "15.00 fps": 15,
    "23.976 fps": 23.976, "24.00 fps": 24, "25.00 fps": 25, 
    "29.97 fps": 29.97, "30.00 fps": 30, "50.00 fps": 50, 
    "59.94 fps": 59.94, "60.00 fps": 60
}

RES_MAP = {
    "1080x1920 (Vertical HD)": (1080, 1920),
    "1920x1080 (Landscape HD)": (1920, 1080),
    "2160x3840 (Vertical 4K)": (2160, 3840),
    "3840x2160 (Landscape 4K)": (3840, 2160),
    "1080x1080 (Square)": (1080, 1080)
}

def tc_to_frames(tc, source_fps, target_fps):
    try:
        clean_tc = tc.replace(';', ':')
        h, m, s, f = list(map(int, clean_tc.split(':')))
        
        # Convert timecode to absolute decimal seconds using the Source FPS
        # This prevents the "rollover" because we read :29 as part of the Source base
        total_seconds = (h * 3600) + (m * 60) + s + (f / source_fps)
        
        # Convert those seconds into frames for the Target XML
        # We use math.floor to match Premiere's internal rounding
        return math.floor(total_seconds * target_fps)
    except:
        return 0

def set_font(run, size=11, bold=False):
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Arial')
    run.font.size = Pt(size)
    run.bold = bold

# --- 2. UI SETUP ---
st.set_page_config(page_title="QOMY Feedback Tool", page_icon="🎬", layout="centered")

st.title("🎬 QOMY Feedback Tool")
st.markdown("Upload your Premiere CSV to generate formatted Feedback Docs and XML Markers.")

# Sidebar Settings
st.sidebar.header("GLOBAL SETTINGS")

# Request: Instruction sentence and Target FPS
st.sidebar.write("Select Premiere Sequence FPS (Target):")
target_fps_choice = st.sidebar.selectbox("Target FPS Dropdown:", list(XML_TIMEBASE_MAP.keys()), index=6, label_visibility="collapsed")
target_fps_val = XML_TIMEBASE_MAP[target_fps_choice]

# Fix: Added Source FPS to prevent rollover
st.sidebar.write("Select Source CSV FPS (Original):")
source_fps_choice = st.sidebar.selectbox("Source FPS Dropdown:", list(XML_TIMEBASE_MAP.keys()), index=7, label_visibility="collapsed")
source_fps_val = XML_TIMEBASE_MAP[source_fps_choice]

# Request: Resolution dropdown
st.sidebar.write("Select Sequence Resolution:")
res_choice = st.sidebar.selectbox("Resolution Dropdown:", list(RES_MAP.keys()), index=0, label_visibility="collapsed")
width, height = RES_MAP[res_choice]

# --- 3. WORKFLOW ---
csv_file = st.file_uploader("Select Premiere CSV", type="csv")
logo_file = st.file_uploader("Upload Logo (Optional)", type=["png", "jpg"])

if csv_file:
    try:
        raw_data = csv_file.read()
        content = ""
        for enc in ['utf-8-sig', 'utf-16', 'cp1252']:
            try:
                content = raw_data.decode(enc)
                if "Marker Name" in content: break
            except: continue
        
        dialect = csv.Sniffer().sniff(content[:2000])
        reader = csv.DictReader(content.splitlines(), dialect=dialect)
        
        doc = Document()
        
        if logo_file:
            doc.add_picture(io.BytesIO(logo_file.read()), width=Inches(1.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Request: Filename logic ([base]feedback)
        base_name = Path(csv_file.name).stem
        final_filename = f"{base_name}feedback"
        
        title_para = doc.add_heading('', 0)
        title_run = title_para.add_run(base_name)
        set_font(title_run, size=18, bold=True)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        xml_markers = ""
        
        for row in reader:
            name = row.get('Marker Name', '').strip()
            desc = row.get('Description', '').strip()
            comment = name if len(name) >= len(desc) else desc
            
            in_tc = row.get('In', '00:00:00:00')
            out_tc = row.get('Out', in_tc)
            ts_display = in_tc if in_tc == out_tc else f"{in_tc} - {out_tc}"
            
            # Word Doc Paragraphs (Arial, Black)
            p_ts = doc.add_paragraph()
            run_ts = p_ts.add_run(ts_display)
            set_font(run_ts, bold=True)

            p_cmt = doc.add_paragraph()
            run_cmt = p_cmt.add_run(comment)
            set_font(run_cmt)
            
            # XML Logic (Corrected Math)
            start_f = tc_to_frames(in_tc, source_fps_val, target_fps_val)
            end_f = tc_to_frames(out_tc, source_fps_val, target_fps_val)
            clean_cmt = comment.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            xml_markers += f"<marker><name>NOTE</name><comment>{clean_cmt}</comment><in>{int(start_f)}</in><out>{int(end_f)}</out></marker>"

        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        # XML Header Configuration
        timebase_str = "24" if target_fps_val == 23.976 else str(int(target_fps_val))
        ntsc_bool = "TRUE" if target_fps_val in [23.976, 29.97, 59.94] else "FALSE"
        
        full_xml = f'<?xml version="1.0" encoding="UTF-8"?><xmeml version="4"><project><children><sequence><name>{final_filename}</name><rate><timebase>{timebase_str}</timebase><ntsc>{ntsc_bool}</ntsc></rate><media><video><format><samplecharacteristics><width>{width}</width><height>{height}</height></samplecharacteristics></format></video></media>{xml_markers}</sequence></children></project></xmeml>'

        st.divider()
        st.success(f"Processed: {base_name}")
        
        c1, c2 = st.columns(2)
        with c1:
            st.download_button("⬇️ Download Docx", data=doc_io, file_name=f"{final_filename}.docx")
        with c2:
            st.download_button("⬇️ Download XML", data=full_xml, file_name=f"{final_filename}.xml")

    except Exception as e:
        st.error(f"Error: {e}")
